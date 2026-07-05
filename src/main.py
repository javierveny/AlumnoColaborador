import os
import io
import json
import re
import shutil
from pathlib import Path
from urllib.parse import quote

import requests
import pdfplumber
from docx import Document
from pydantic import BaseModel
from openai import OpenAI

from fastapi import FastAPI, UploadFile, File
from fastapi.responses import StreamingResponse

from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from minio import Minio

# ==========================================
# 1. CONFIGURACIÓN Y VARIABLES DE ENTORNO
# ==========================================
uri_mongo = os.getenv("MONGO_URI")
ollama_host = os.getenv("OLLAMA_HOST", "http://ollama:11434")
api_key_openrouter = os.getenv("OPENROUTER_API_KEY")

minio_endpoint = os.getenv("MINIO_ENDPOINT", "minio:9000")
minio_access_key = os.getenv("MINIO_ROOT_USER")
minio_secret_key = os.getenv("MINIO_ROOT_PASSWORD")
minio_bucket = os.getenv("MINIO_BUCKET", "documents")

# Variables globales para el procesamiento en memoria
texto_del_documento = ""
texto_metadatos = ""
url_kahoot = "https://create.kahoot.it/rest/kahoots/"

# ==========================================
# 2. INICIALIZACIÓN DE CLIENTES EXTERNOS
# ==========================================
# Cliente MinIO
cliente_minio = Minio(
    minio_endpoint,
    access_key=minio_access_key,
    secret_key=minio_secret_key,
    secure=False
)

# Cliente OpenAI (Conectado a Ollama en local)
modelo_llm = os.getenv("OLLAMA_MODEL", "llama3.2")
cliente = OpenAI(
    base_url=f"{ollama_host}/v1",
    api_key="ollama",
)

# ==========================================
# 3. DEFINICIÓN DE LA API (FastAPI)
# ==========================================
app = FastAPI(
    title="API Proyecto Alumno Colaborador",
    description="API para procesar y consultar exámenes extraídos con IA local y almacenamiento distribuido.",
    version="1.0.0"
)

class PeticionKahoot(BaseModel):
    url: str

@app.get("/preguntas", tags=["Consultas"])
def obtener_preguntas(asignatura: str = None, curso: str = None, nivel_bloom: str = None):
    """Busca preguntas en MongoDB aplicando filtros dinámicos."""
    clienteMongo = MongoClient(uri_mongo, server_api=ServerApi('1'))
    db = clienteMongo["proyecto_alumno_colaborador"]
    coleccion = db["preguntas_examenes"]

    filtro = {}
    if asignatura:
        filtro["Asignatura"] = {"$regex": asignatura, "$options": "i"}
    if curso:
        filtro["Curso"] = curso
    if nivel_bloom:
        filtro["Nivel_cognitivo_Bloom"] = {"$regex": nivel_bloom, "$options": "i"}

    preguntas_db = list(coleccion.find(filtro, {"_id": 0}))

    return {
        "filtros_aplicados": filtro,
        "total_preguntas": len(preguntas_db),
        "datos": preguntas_db
    }

@app.get("/documentos/{nombre_archivo}", tags=["Almacenamiento"])
def descargar_documento(nombre_archivo: str):
    """Descarga un documento original almacenado en MinIO."""
    try:
        respuesta_minio = cliente_minio.get_object(minio_bucket, nombre_archivo)
        
        def iterar_archivo():
            try:
                for pedazo in respuesta_minio.stream(32 * 1024):
                    yield pedazo
            finally:
                respuesta_minio.close()
                respuesta_minio.release_conn()

        nombre_codificado = quote(nombre_archivo)
        return StreamingResponse(
            iterar_archivo(), 
            media_type="application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{nombre_codificado}"}
        )
    except Exception as e:
        return {"error": f"No se pudo descargar el archivo: {str(e)}"}

@app.post("/procesar-kahoot", tags=["Procesamiento IA"])
def procesar_kahoot(peticion: PeticionKahoot):
    """Procesa un JSON de Kahoot mediante su URL, extrae preguntas con IA y las guarda en BD."""
    global texto_del_documento
    link_original = peticion.url
    
    try:
        id_kahoot = Path(link_original).name
        nombre_en_minio = f"link_kahoot_{id_kahoot}.txt"
        subir_link_a_minio(link_original, nombre_en_minio)

        linkKahoot = url_kahoot + str(id_kahoot)
        respuesta = requests.get(linkKahoot)
        
        if respuesta.status_code != 200:
            return {"error": "No se pudo descargar el Kahoot. Comprueba el link."}
            
        texto_del_documento = respuesta.text
        llamada_llm(info=False)

        return {
            "estado": "Éxito",
            "mensaje": f"El Kahoot {id_kahoot} ha sido procesado y guardado en MongoDB."
        }
    except Exception as e:
        return {"error": f"Error inesperado: {str(e)}"}

@app.post("/procesar-documento", tags=["Procesamiento IA"])
def procesar_documento(examen: UploadFile = File(...), archivo_metadatos: UploadFile = File(None)):
    """Procesa un PDF/DOCX, lo sube a MinIO y extrae su contenido usando LLMs."""
    global texto_del_documento
    global texto_metadatos

    try:
        ruta_examen = f"src/{examen.filename}"
        with open(ruta_examen, "wb") as buffer:
            shutil.copyfileobj(examen.file, buffer)

        subir_a_minio(ruta_examen)
        ruta_fichero = Path(ruta_examen)

        if ruta_fichero.suffix == ".docx":
            texto_del_documento = extraer_texto_docx(ruta_fichero)
        elif ruta_fichero.suffix == ".pdf":
            texto_del_documento = extraer_texto_pdf(ruta_fichero)
        else:
            os.remove(ruta_examen)
            return {"error": "Formato no soportado. Solo .pdf o .docx"}

        tiene_metadatos = False
        if archivo_metadatos:
            ruta_meta = f"src/{archivo_metadatos.filename}"
            with open(ruta_meta, "wb") as buffer:
                shutil.copyfileobj(archivo_metadatos.file, buffer)
            
            ruta_meta_path = Path(ruta_meta)
            if ruta_meta_path.suffix == ".docx":
                texto_metadatos = extraer_texto_docx(ruta_meta_path)
                tiene_metadatos = True
            
            os.remove(ruta_meta)

        llamada_llm(info=tiene_metadatos)
        os.remove(ruta_examen)

        return {
            "estado": "Éxito",
            "mensaje": f"El documento {examen.filename} ha sido procesado correctamente."
        }
    except Exception as e:
        return {"error": f"Error inesperado: {str(e)}"}

# ==========================================
# 4. FUNCIONES AUXILIARES Y DE EXTRACCIÓN
# ==========================================
def subir_a_minio(ruta_archivo):
    nombre_archivo = Path(ruta_archivo).name
    if not cliente_minio.bucket_exists(minio_bucket):
        cliente_minio.make_bucket(minio_bucket)
    try:
        cliente_minio.fput_object(minio_bucket, nombre_archivo, ruta_archivo)
        return nombre_archivo
    except Exception as err:
        print(f"❌ Error subiendo a MinIO: {err}")
        return None

def subir_link_a_minio(link, nombre_archivo):
    datos_bytes = link.encode('utf-8')
    flujo_datos = io.BytesIO(datos_bytes)
    if not cliente_minio.bucket_exists(minio_bucket):
        cliente_minio.make_bucket(minio_bucket)
    try:
        cliente_minio.put_object(minio_bucket, nombre_archivo, flujo_datos, length=len(datos_bytes), content_type='text/plain')
    except Exception as err:
        print(f"❌ Error guardando link en MinIO: {err}")

def extraer_texto_docx(ruta_docx):
    doc = Document(ruta_docx)
    output = []
    for element in doc.element.body:
        if element.tag.endswith('p'):
            paragraph = [p for p in doc.paragraphs if p._element == element]
            if paragraph and paragraph[0].text.strip():
                output.append(paragraph[0].text.strip())
        elif element.tag.endswith('tbl'):
            tabla_obj = [t for t in doc.tables if t._element == element][0]
            output.append("\n[TABLA_START]")
            for i, row in enumerate(tabla_obj.rows):
                celdas = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                output.append("| " + " | ".join(celdas) + " |")
                if i == 0:
                    separador = "| " + " | ".join(["---"] * len(celdas)) + " |"
                    output.append(separador)
            output.append("[TABLA_END]\n")
    return "\n".join(output)

def extraer_texto_pdf(ruta_pdf):
    output = []
    with pdfplumber.open(ruta_pdf) as pdf:
        for pagina in pdf.pages:
            texto_pagina = pagina.extract_text(x_tolerance=3, y_tolerance=3, layout=False)
            if texto_pagina:
                output.append(texto_pagina)
            tablas = pagina.extract_tables()
            for tabla in tablas:
                output.append("\n[TABLA_START]")
                for fila in tabla:
                    fila_limpia = [str(celda).replace('\n', ' ') if celda else "" for celda in fila]
                    output.append("| " + " | ".join(fila_limpia) + " |")
                output.append("[TABLA_END]\n")
    return "\n".join(output)

# ==========================================
# 5. FUNCIONES DE LLM Y BASE DE DATOS
# ==========================================
def procesar_respuesta_llm(respuesta_llm):
    texto_limpio = respuesta_llm.strip()
    if texto_limpio.startswith("```"):
        lineas = texto_limpio.splitlines()
        texto_limpio = "\n".join(lineas[1:-1]) if lineas[0].startswith("```") else texto_limpio
    try:
        return json.loads(texto_limpio)
    except json.JSONDecodeError as e:
        print(f"Error al decodificar: {e}")
        return None

def comprobar_respuesta_llm(respuesta_llm, campos_requeridos):
    # Uso de .*? (lazy) para evitar capturar texto basura generado por el LLM tras el JSON
    match = re.search(r"\[.*?\]", respuesta_llm, re.DOTALL)
    if not match:
        print("No es un JSON válido")
        return None
    try:
        datos_json = json.loads(match.group(0))
    except json.JSONDecodeError:
        print("Había algun error en el JSON")
        return None
    if len(datos_json) == 0:
        print("❌ El JSON está vacío.")
        return None
    for pregunta in datos_json:
        if set(pregunta.keys()) != campos_requeridos:
            print(f"Error de formato: La IA ha devuelto campos incorrectos.")
            return None
    return datos_json

def llamada_llm(info=False):
    i = 5
    datos = None
    while i > 0:
        if not info:
            completion = cliente.chat.completions.create(
                model=modelo_llm,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Eres un extractor de datos técnicos preciso. Tu única función es transformar un JSON de Kahoot a una LISTA de objetos JSON.\n"
                            "REGLA DE ORO: Debes generar UN objeto JSON por cada pregunta. No omitas ninguna.\n"
                            "REGLA DE COPIA: El 'Enunciado_completo' debe incluir la pregunta seguida de todas sus opciones.\n"
                            "REGLA DE CURSO: Elige estrictamente entre 'Primero', 'Segundo', 'Tercero', 'Cuarto'.\n"
                            "REGLA DE SOLUCIÓN: Escribe el texto exacto de la opción marcada como 'correct': true."
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            "Transforma TODAS las preguntas de este examen a este formato de lista de Python [{}, {}, ...]:\n"
                            "[\n  {\n"
                            '    "Asignatura": "...",\n'
                            '    "Curso": "Primero, Segundo, Tercero o Cuarto",\n'
                            '    "Estudios": "...",\n'
                            '    "Enunciado_completo": "...",\n'
                            '    "Solución": "..."\n'
                            "  }\n]\n\n"
                            f"TEXTO DEL EXAMEN:\n{texto_del_documento}"
                        )
                    }
                ],
                temperature=0
            )
            campos_requeridos = {"Asignatura", "Curso", "Estudios", "Enunciado_completo", "Solución"}
        else:
            completion = cliente.chat.completions.create(
                model=modelo_llm,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Eres un extractor de datos académicos estructurados. Tu salida debe ser EXCLUSIVAMENTE una LISTA de objetos JSON.\n"
                            "REGLA CRÍTICA DE COPIA: Copia EXACTAMENTE los campos de los METADATOS para todas las preguntas.\n"
                            "REGLA CRÍTICA DE EXTRACCIÓN: Extrae 'Enunciado_completo' y 'Solución' individualmente del texto del EXAMEN.\n"
                            "REGLA CRÍTICA PARA 'Curso': Elige únicamente entre: 'Primero', 'Segundo', 'Tercero', 'Cuarto'.\n"
                            "Si hay preguntas tipo test, pon todas las posibles respuestas en el enunciado."
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            "Analiza los metadatos y el texto del examen, y genera la lista de JSONs con este formato:\n"
                            "[\n  {\n"
                            '    "Asignatura": "...",\n'
                            '    "Curso": "...",\n'
                            '    "Estudios": "...",\n'
                            '    "Autores": "...",\n'
                            '    "Nivel_cognitivo_Bloom": "...",\n'
                            '    "Tipo_pregunta": "...",\n'
                            '    "Competencias_relacionadas": "...",\n'
                            '    "Nivel de dificultad": "...",\n'
                            '    "Tema / topic": "...",\n'
                            '    "Idioma": "...",\n'
                            '    "Enunciado_completo": "...",\n'
                            '    "Solución": "..."\n'
                            "  }\n]\n\n"
                            f"--- METADATOS GLOBALES ---\n{texto_metadatos}\n\n"
                            f"--- TEXTO DEL EXAMEN ---\n{texto_del_documento}"
                        )
                    }
                ]
            )
            campos_requeridos = {"Asignatura", "Curso", "Estudios", "Autores", "Nivel_cognitivo_Bloom", "Tipo_pregunta", "Competencias_relacionadas", "Nivel de dificultad", "Tema / topic", "Idioma", "Enunciado_completo", "Solución"}

        respuesta_llm = completion.choices[0].message.content
        datos = comprobar_respuesta_llm(respuesta_llm, campos_requeridos)

        if datos is None:
            i -= 1
        else:
            print("🎉 ¡La IA ha acertado el formato!")
            break 

    if datos is None:
        print("⚠️ Se han agotado los 5 intentos. Abortando esta extracción.")
        return
    else:
        with open("src/data.json", "w", encoding="utf-8") as f:
            json.dump(datos, f, ensure_ascii=False, indent=4)
        segunda_iteracion_llm(datos)

def segunda_iteracion_llm(datos):
    with open("src/competencias.json", 'r', encoding='utf-8') as file:
        competencias = json.load(file)

    clienteMongo = MongoClient(uri_mongo, server_api=ServerApi('1'))
    coleccion = clienteMongo["proyecto_alumno_colaborador"]["preguntas_examenes"]

    for elemento in datos:
        completion = cliente.chat.completions.create(
            model=modelo_llm,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Eres un clasificador académico estricto. Tu tarea es analizar una pregunta y devolver EXCLUSIVAMENTE un único objeto JSON.\n"
                        "1. 'Nivel_cognitivo_Bloom': 'Recordar', 'Entender', 'Aplicar', 'Analizar', 'Evaluar', 'Crear'.\n"
                        "2. 'Tipo_pregunta': 'Resolución de problemas', 'Diseño/Modelado', 'Análisis de caso práctico', 'Test. Opción múltiple', 'Respuesta corta', 'Codificación', 'Ensayo breve'.\n"
                        "3. 'Competencias_detalladas': Busca los códigos en el CATÁLOGO y devuelve un array con las descripciones."
                    )
                },
                {
                    "role": "user",
                    "content": (
                        "Genera el JSON con la clasificación. Devuelve SOLO las 3 claves:\n"
                        "{\n"
                        '  "Nivel_cognitivo_Bloom": "...",\n'
                        '  "Tipo_pregunta": "...",\n'
                        '  "Competencias_detalladas": ["..."]\n'
                        "}\n\n"
                        f"--- CATÁLOGO ---\n{competencias}\n\n"
                        f"--- PREGUNTA ---\n{json.dumps(elemento, ensure_ascii=False)}"
                    )
                }
            ]
        )
        respuesta_llm = procesar_respuesta_llm(completion.choices[0].message.content)

        if respuesta_llm and len(respuesta_llm) > 0:
            elemento.update(respuesta_llm)
            coleccion.insert_one(elemento)
            print(f"✅ Guardado en MongoDB: {elemento.get('Asignatura')}")
        else:
            print("⚠️ El LLM no devolvió datos válidos para esta pregunta.")
